# selenium 4
from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service as ChromiumService
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.core.os_manager import ChromeType
import time
import docx
from PyQt5 import *
from status_ui import *  # импорт нашего сгенерированного файла
import sys
from colorama import *
from PyQt5.QtWidgets import QFileDialog
import os
from tkinter import ttk
import tkinter
from tkinter import *
from tkinter.messagebox import showerror, showwarning, showinfo
init()
from datetime import datetime
now = datetime.now()
start_time = time.time()
class mywindow(QtWidgets.QMainWindow):
    
    def __init__(self):
        super(mywindow, self).__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.ui.StatBtn_2.clicked.connect(self.test)
        self.ui.StatBtn_3.clicked.connect(self.pass_mng)
        self.ui.pushButton.clicked.connect(self.open)
        global myfile
        global g_file

        #self.ui.progressBar.setMinimum(0)
        #self.ui.progressBar.setMaximum(lines)
    def open(self):
        print()
        global g_file
        myfile = QFileDialog.getOpenFileName()
        print(str(myfile[0]))
        g_file = str(myfile[0])
    def test(self):

        dir_path = os.path.dirname(g_file)
        print(dir_path)
        print('StatCheckBeta') #######################ИМЯ ФАЙЛА########################
        showinfo(title=None, message="Проверка начата. НЕ Закрывайте окно даже, если кажется, что программа зависла")
        doc = docx.Document()
        doc.add_heading('Статусы на '+str(now.day)+'.'+str(now.month)+'.'+str(now.year))
        with open(g_file, "r", encoding='utf-8') as file:
          
          current_line_num = 0
          num_of_err = 0
          for item in file:
                #self.ui.progressBar.setValue(current_line_num+ 1)
                p_list = item.split()
                print(" ")
                print("==========================================")
                print(" ")
                print("L: "+ p_list[0]+" P: "+p_list[1])
                options_of_driver = webdriver.ChromeOptions()
                options_of_driver.add_argument("--headless")
                options_of_driver.add_argument("--disable-blink-features=AutomationControlled")
                options_of_driver.add_experimental_option("prefs", {"profile.managed_default_content_settings.images": 2})
                driver = webdriver.Chrome(service=ChromiumService(ChromeDriverManager(chrome_type=ChromeType.CHROMIUM).install()), options=options_of_driver)
                try:
                  driver.get("https://портал-тп.рф/platform/portal/tehprisEE_profileMyMsg")
                  driver.implicitly_wait(5)
                  email_int = driver.find_element(By.ID, "workplaceTopForm:j_mail_login")
                  email_int.send_keys(p_list[0])
                  passw_int = driver.find_element(By.ID, "workplaceTopForm:j_password")
                  passw_int.send_keys(p_list[1])
                  driver.find_element(By.ID, "workplaceTopForm:loginBtn").click()
                  driver.implicitly_wait(8) # Даем время на загрузку сайта
                #Окончание входа
                  print("Приступаю к проверке")
                #Считывание данных
                  
                  num_of_stat = 0
                  while num_of_stat <= 5:
                    driver.implicitly_wait(2)
                    print("Заявка №" + str(num_of_stat))
                    try:

                      zayavka_class = driver.find_element(By.CLASS_NAME, "item-message")
                      #print("получен класс")
                      zayavka_id = zayavka_class.get_attribute("id")
                      #s = "workplaceForm:messagesProfile:myRARProfile:requestsAll:1:j_idt1386"
                      new_s = zayavka_id[:55] + str(num_of_stat)  + zayavka_id[56:]
                      print(Fore.RED+ new_s+ Fore.RESET)
                      zayavka = driver.find_element(By.ID, new_s)
                  #Считывание данных
                      #print("Заявка №" + str(num_of_stat) +" найдена")
                      status = zayavka.find_element(By.CLASS_NAME,"answer-status-in-list").text
                      status_num = zayavka.find_element(By.CLASS_NAME,"item-message__next-id").text
                      fio = driver.find_element(By.ID, "workplaceForm:fioLabel").text
                      date = zayavka.find_element(By.CLASS_NAME,"item-message__option-value").text
                      itog_stat = fio +" "+ status_num +" "+date +" "+ status
                      print(Fore.GREEN + itog_stat+ Fore.RESET)
                      num_of_stat= num_of_stat+1
                      stat_to_ui = [itog_stat]
                      self.ui.Individ.addItems(stat_to_ui)
                      num_of_err = 0
                      doc.add_paragraph(itog_stat)
                    except:
                      num_of_stat= 6
                      print("Заявок больше нет")
                      print(num_of_err)
                      driver.quit()
                      num_of_err = num_of_err+1
                      if num_of_err == 2:

                          
                          showwarning(title='Предупреждение', message="Рекомендуется проверить соединение с интернетом, правильность пароля или обновить программу",)
                          err_msg = "[ERR] ВНИМАНИЕ!!!! Пользователь с данными "+"L: "+ p_list[0]+" P: "+p_list[1] +" не был проверен"
                  #tkinter.showwarning(title='Предупреждение', message="Рекомендуется проверить соединение с интернетом, правильность пароля или обновить программу",)
                          doc.add_paragraph(err_msg)
                          err_to_ui = [err_msg]
                          self.ui.Individ.addItems(err_to_ui)
                          print(err_msg)

                      else:
                        print("Колво ошибок: "+ str(num_of_err))
                  

                except:
                  err_msg = "[ERR] ВНИМАНИЕ!!!! Пользователь с данными "+"L: "+ p_list[0]+" P: "+p_list[1] +" не был проверен"
                  #tkinter.showwarning(title='Предупреждение', message="Рекомендуется проверить соединение с интернетом, правильность пароля или обновить программу",)
                  doc.add_paragraph(err_msg)
                  err_to_ui = [err_msg]
                  self.ui.Individ.addItems(err_to_ui)
                  print(Fore.RED + err_msg + Fore.RESET)

                  driver.quit()
          
          
        doc.save(dir_path+'/Статусы.docx')
        #tkinter.showwarning(title='Предупреждение', message="Рекомендуется проверить соединение с интернетом, правильность пароля или обновить программу",)
        print("проверка завершена")
        print("--- %s seconds ---" % (time.time() - start_time))
    def pass_mng(self):
      os.system("passmng.exe")
      
      

            
app = QtWidgets.QApplication([])
application = mywindow()
application.show()
 
sys.exit(app.exec())
